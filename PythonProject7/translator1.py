from tkinter import *
from tkinter import ttk, messagebox, filedialog
import googletrans
from googletrans import Translator
import docx
import PyPDF2

def change_label():
    c1 = one_combo.get()
    c2 = two_combo.get()
    label_1.configure(text=c1.upper())
    label_2.configure(text=c2.upper())

def translate_text(text, src_lang, dest_lang):
    translator = Translator()
    translated = translator.translate(text, src=src_lang, dest=dest_lang)
    return translated.text

def current_trans():
    text_ = text_1.get(1.0, END)
    t1 = Translator()
    try:
        trans_text = t1.translate(text_, src=one_combo.get(), dest=two_combo.get()).text
        text_2.delete(1.0, END)
        text_2.insert(END, trans_text)
    except Exception as e:
        messagebox.showerror("Translation Error", str(e))

def load_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("Text files", "*.txt"), ("Word documents", "*.docx"), ("PDF files", "*.pdf")]
    )
    if file_path:
        try:
            content = ""
            if file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as file:
                    content = file.read()
            elif file_path.endswith(".docx"):
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            elif file_path.endswith(".pdf"):
                with open(file_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        content += page.extract_text()

            text_1.delete(1.0, END)
            text_1.insert(END, content)

        except Exception as e:
            messagebox.showerror("File Error", f"Failed to read file: {e}")

def save_translated_file():
    trans_text = text_2.get(1.0, END).strip()
    if not trans_text:
        messagebox.showwarning("No Translation", "Please translate text before saving.")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".txt",
        filetypes=[
            ("Text file", "*.txt"),
            ("Word document", "*.docx")
        ],
    )
    if file_path:
        try:
            if file_path.endswith(".txt"):
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(trans_text)
            elif file_path.endswith(".docx"):
                doc = docx.Document()
                doc.add_paragraph(trans_text)
                doc.save(file_path)

            messagebox.showinfo("Success", "File saved successfully!")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save file: {e}")

window = Tk()
window.title("LANGUAGE TRANSLATOR")
window.geometry("1080x520")
window.config(background="light pink")

trans_language = googletrans.LANGUAGES
lang_a = list(trans_language.values())

one_combo = ttk.Combobox(window, values=lang_a, font="Arial 15 bold", state="r")
one_combo.place(x=110, y=20)
one_combo.set("English")

two_combo = ttk.Combobox(window, values=lang_a, font="Arial 15 bold", state="r")
two_combo.place(x=730, y=20)
two_combo.set("Select Language")

label_1 = Label(window, text="ENGLISH", font="Arial 31 bold", bg="light green", width=17, bd=4, relief=GROOVE)
label_1.place(x=10, y=50)

label_2 = Label(window, text="LANGUAGE", font="Arial 31 bold", bg="light green", width=17, bd=4, relief=GROOVE)
label_2.place(x=600, y=50)

f_1 = Frame(window, bg="black", bd=5)
f_1.place(x=10, y=118, width=440, height=250)
text_1 = Text(f_1, font="Arial 14", bg="light blue", relief=GROOVE, wrap=WORD)
text_1.pack(fill=BOTH, expand=True)
scrollbar_one = Scrollbar(f_1, command=text_1.yview)
scrollbar_one.pack(side="right", fill="y")
text_1.config(yscrollcommand=scrollbar_one.set)

f_2 = Frame(window, bg="black", bd=5)
f_2.place(x=595, y=118, width=440, height=250)
text_2 = Text(f_2, font="Arial 14", bg="light blue", relief=GROOVE, wrap=WORD)
text_2.pack(fill=BOTH, expand=True)
scrollbar_two = Scrollbar(f_2, command=text_2.yview)
scrollbar_two.pack(side="right", fill="y")
text_2.config(yscrollcommand=scrollbar_two.set)

button_translate = Button(window, text="TRANSLATE", font="Arial 16 bold italic", activebackground="red",
                          bg="red", cursor="hand2", bd=4, fg="white", command=current_trans)
button_translate.place(x=450, y=400)

button_load = Button(window, text="LOAD FILE", font="Arial 14 bold", bg="blue", fg="white", command=load_file)
button_load.place(x=450, y=350)

button_save = Button(window, text="SAVE FILE", font="Arial 14 bold", bg="green", fg="white", command=save_translated_file)
button_save.place(x=450, y=450)

change_label()
window.mainloop()